"""LlamaIndex document loader registry wiring local workspaces and cloud loaders."""

from __future__ import annotations

import logging
import mimetypes
from dataclasses import dataclass
from email.parser import BytesParser
from email.policy import default as default_email_policy
from importlib import import_module
from importlib.util import find_spec
from pathlib import Path
from typing import Callable, Dict, Iterable, List, Optional, Protocol, Tuple

from .fallback import FallbackDocument, MetadataModeEnum

from backend.app.models.api import IngestionSource
from backend.app.utils.text import read_text

from .ocr import OcrEngine, OcrResult
from .settings import LlamaIndexRuntimeConfig
from .utils import compute_sha256


@dataclass
class DocumentLike(Protocol):
    metadata: Dict[str, object] | None

    def get_content(self, metadata_mode: object | None = None) -> str:  # pragma: no cover - protocol
        ...


def _has_spec(path: str) -> bool:
    try:
        return find_spec(path) is not None
    except ModuleNotFoundError:
        return False


def _resolve_llama_index_document() -> type[DocumentLike]:
    if not _has_spec("llama_index.core"):
        return FallbackDocument
    try:
        module = import_module("llama_index.core")
        return getattr(module, "Document")
    except (ModuleNotFoundError, AttributeError):
        return FallbackDocument


def _resolve_llama_index_metadata_mode() -> object:
    if not _has_spec("llama_index.core.schema"):
        return MetadataModeEnum
    try:
        module = import_module("llama_index.core.schema")
        return getattr(module, "MetadataMode")
    except (ModuleNotFoundError, AttributeError):
        return MetadataModeEnum


Document = _resolve_llama_index_document()
MetadataMode = _resolve_llama_index_metadata_mode()
METADATA_MODE_ALL = getattr(MetadataMode, "ALL", MetadataModeEnum.ALL)


@dataclass
class LoadedDocument:
    """Container holding a LlamaIndex document plus provenance metadata."""

    source: IngestionSource
    path: Path
    document: DocumentLike
    text: str
    checksum: str
    metadata: Dict[str, object]
    ocr: Optional[OcrResult]


class HubLoaderFactory:
    """Instantiate LlamaHub loaders on demand."""

    def __call__(self, loader_name: str):
        if not _has_spec("llama_hub.utils"):
            raise RuntimeError(
                "llama-hub is required to load remote ingestion sources. Install llama-hub to "
                "enable SharePoint, OneDrive, Gmail, IMAP, and Google Drive connectors."
            )
        module = import_module("llama_hub.utils")
        download_loader = getattr(module, "download_loader")
        return download_loader(loader_name)


class LoaderRegistry:
    """Bridge between storage connectors and LlamaIndex documents."""

    _EMAIL_EXTENSIONS = {".eml", ".msg"}
    _OCR_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff"}
    _DOCX_EXTENSIONS = {".docx"}
    _PDF_EXTENSION = ".pdf"

    def __init__(
        self,
        runtime_config: LlamaIndexRuntimeConfig,
        ocr_engine: OcrEngine,
        *,
        logger: logging.Logger,
        credential_resolver: Callable[[str], Dict[str, str]] | None = None,
        hub_factory: HubLoaderFactory | None = None,
    ) -> None:
        self.runtime_config = runtime_config
        self.ocr_engine = ocr_engine
        self.logger = logger
        self.hub_factory = hub_factory or HubLoaderFactory()
        self._resolve_credentials = credential_resolver

    def load_documents(
        self, materialized_root: Path, source: IngestionSource, *, origin: str
    ) -> List[LoadedDocument]:
        source_type = source.type.lower()
        if source_type in {"sharepoint", "onedrive", "gmail", "imap", "gdrive"}:
            return self._load_via_llamahub(source, origin)
        return list(self._load_from_workspace(materialized_root, source, origin))

    # ------------------------------------------------------------------
    def _load_from_workspace(
        self, root: Path, source: IngestionSource, origin: str
    ) -> Iterable[LoadedDocument]:
        for path in sorted(root.rglob("*")):
            if not path.is_file():
                continue
            suffix = path.suffix.lower()
            if suffix == self._PDF_EXTENSION:
                yield self._load_pdf(path, source, origin)
            elif suffix in self._OCR_IMAGE_EXTENSIONS:
                yield self._load_image(path, source, origin)
            elif suffix in self._EMAIL_EXTENSIONS:
                yield self._load_email(path, source, origin)
            elif suffix in self._DOCX_EXTENSIONS:
                yield self._load_docx(path, source, origin)
            else:
                yield self._load_text(path, source, origin)

    def _load_text(self, path: Path, source: IngestionSource, origin: str) -> LoadedDocument:
        text = read_text(path)
        metadata = self._base_metadata(path, source, origin)
        document = Document(text=text, metadata=metadata, metadata_mode=METADATA_MODE_ALL)
        checksum = compute_sha256(path)
        return LoadedDocument(source=source, path=path, document=document, text=text, checksum=checksum, metadata=metadata, ocr=None)

    def _load_pdf(self, path: Path, source: IngestionSource, origin: str) -> LoadedDocument:
        ocr_result = self.ocr_engine.extract_from_pdf(path)
        text = ocr_result.text or read_text(path)
        metadata = self._base_metadata(path, source, origin)
        metadata.update(
            {
                "ocr_engine": ocr_result.engine,
                "ocr_confidence": ocr_result.confidence,
                "ocr_tokens": ocr_result.tokens,
            }
        )
        document = Document(text=text, metadata=metadata, metadata_mode=METADATA_MODE_ALL)
        checksum = compute_sha256(path)
        return LoadedDocument(source=source, path=path, document=document, text=text, checksum=checksum, metadata=metadata, ocr=ocr_result)

    def _load_image(self, path: Path, source: IngestionSource, origin: str) -> LoadedDocument:
        ocr_result = self.ocr_engine.extract_from_image(path)
        metadata = self._base_metadata(path, source, origin)
        metadata.update(
            {
                "ocr_engine": ocr_result.engine,
                "ocr_confidence": ocr_result.confidence,
                "ocr_tokens": ocr_result.tokens,
            }
        )
        document = Document(text=ocr_result.text, metadata=metadata, metadata_mode=METADATA_MODE_ALL)
        checksum = compute_sha256(path)
        return LoadedDocument(source=source, path=path, document=document, text=ocr_result.text, checksum=checksum, metadata=metadata, ocr=ocr_result)

    def _load_docx(self, path: Path, source: IngestionSource, origin: str) -> LoadedDocument:
        from docx import Document as DocxDocument

        docx = DocxDocument(str(path))
        text = "\n".join(paragraph.text for paragraph in docx.paragraphs)
        metadata = self._base_metadata(path, source, origin)
        document = Document(text=text, metadata=metadata, metadata_mode=METADATA_MODE_ALL)
        checksum = compute_sha256(path)
        return LoadedDocument(source=source, path=path, document=document, text=text, checksum=checksum, metadata=metadata, ocr=None)

    def _load_email(self, path: Path, source: IngestionSource, origin: str) -> LoadedDocument:
        raw = path.read_bytes()
        if path.suffix.lower() == ".msg":
            from extract_msg import Message  # type: ignore

            message = Message(path)
            text = message.body or ""
            metadata = {
                "subject": message.subject,
                "sender": message.sender,
                "recipients": message.to,
                "date": message.date,
            }
        else:
            parsed = BytesParser(policy=default_email_policy).parsebytes(raw)
            text = parsed.get_body(preferencelist=("plain",)).get_content() if parsed.get_body() else parsed.get_payload()
            metadata = {
                "subject": parsed.get("Subject"),
                "sender": parsed.get("From"),
                "recipients": parsed.get_all("To"),
                "date": parsed.get("Date"),
            }
        base = self._base_metadata(path, source, origin)
        base.update({f"email_{key}": value for key, value in metadata.items() if value})
        document = Document(text=text, metadata=base, metadata_mode=MetadataMode.ALL)
        checksum = compute_sha256(path)
        return LoadedDocument(source=source, path=path, document=document, text=text, checksum=checksum, metadata=base, ocr=None)

    def _load_via_llamahub(self, source: IngestionSource, origin: str) -> List[LoadedDocument]:
        loader_map = {
            "sharepoint": "SharePointReader",
            "onedrive": "OneDriveReader",
            "gmail": "GmailReader",
            "imap": "IMAPReader",
            "gdrive": "GoogleDriveReader",
        }
        loader_name = loader_map[source.type.lower()]
        loader_cls = self.hub_factory(loader_name)
        credentials = self._credentials_for(source)
        init_args, load_args = self._prepare_loader_args(source, credentials)
        loader = loader_cls(**init_args)
        documents: List[LoadedDocument] = []
        for doc in loader.load_data(**load_args):
            text = doc.get_content(metadata_mode=METADATA_MODE_ALL)
            metadata = dict(doc.metadata or {})
            metadata.update({"origin_uri": origin, "source_type": source.type.lower()})
            fake_path = Path(metadata.get("path", source.path or "remote"))
            checksum = compute_sha256(text.encode("utf-8"))
            documents.append(
                LoadedDocument(
                    source=source,
                    path=fake_path,
                    document=Document(text=text, metadata=metadata, metadata_mode=METADATA_MODE_ALL),
                    text=text,
                    checksum=checksum,
                    metadata=metadata,
                    ocr=None,
                )
            )
        return documents

    def _base_metadata(self, path: Path, source: IngestionSource, origin: str) -> Dict[str, object]:
        mime_type, _ = mimetypes.guess_type(path.name)
        return {
            "origin_uri": origin,
            "source_type": source.type.lower(),
            "file_name": path.name,
            "mime_type": mime_type or "application/octet-stream",
            "size_bytes": path.stat().st_size,
        }

    def _credentials_for(self, source: IngestionSource) -> Dict[str, str]:
        if not source.credRef:
            return {}
        if not self._resolve_credentials:
            raise RuntimeError("Credential resolver not configured for LoaderRegistry")
        payload = self._resolve_credentials(source.credRef)
        return {key: str(value) for key, value in payload.items()}

    def _prepare_loader_args(
        self, source: IngestionSource, credentials: Dict[str, str]
    ) -> Tuple[Dict[str, str], Dict[str, str]]:
        source_type = source.type.lower()
        init_args: Dict[str, str] = {}
        load_args: Dict[str, str] = {}

        if source_type == "sharepoint":
            init_args = {
                key: credentials[key]
                for key in ("client_id", "client_secret", "tenant_id", "site_url")
                if key in credentials
            }
            load_args = {
                key: value
                for key, value in {
                    "document_library": credentials.get("document_library"),
                    "relative_path": source.path or credentials.get("relative_path"),
                }.items()
                if value is not None
            }
        elif source_type == "onedrive":
            init_args = {
                key: credentials[key]
                for key in ("client_id", "client_secret", "tenant_id")
                if key in credentials
            }
            load_args = {
                key: value
                for key, value in {
                    "drive_id": credentials.get("drive_id"),
                    "folder_path": source.path or credentials.get("folder_path"),
                }.items()
                if value is not None
            }
        elif source_type == "gmail":
            init_args = {
                key: credentials[key]
                for key in ("token_path", "credentials_path", "user_id")
                if key in credentials
            }
            load_args = {"query": source.path or credentials.get("query", "in:inbox")}
        elif source_type == "imap":
            init_args = {
                key: credentials[key]
                for key in ("host", "user", "password", "ssl")
                if key in credentials
            }
            load_args = {"mailbox": source.path or credentials.get("mailbox", "INBOX")}
        elif source_type == "gdrive":
            init_args = {
                key: credentials[key]
                for key in ("service_account_key", "client_id", "client_secret")
                if key in credentials
            }
            load_args = {
                key: value
                for key, value in {
                    "folder_id": credentials.get("folder_id"),
                    "file_ids": credentials.get("file_ids"),
                }.items()
                if value is not None
            }
        else:
            load_args = {}

        return init_args, load_args


__all__ = ["LoaderRegistry", "LoadedDocument", "HubLoaderFactory"]
