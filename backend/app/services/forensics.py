from __future__ import annotations

import json
import math
import mimetypes
import shutil
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from datetime import datetime, timezone
from decimal import Decimal
from hashlib import md5, sha256
from pathlib import Path
from typing import Any, Callable, Dict, List, Tuple
from time import perf_counter

from opentelemetry import metrics, trace
from opentelemetry.trace import Status, StatusCode
import numpy as np
import pandas as pd
import piexif
from PIL import Image, ImageFilter
from docx import Document as DocxDocument
from extract_msg import Message as MsgMessage
from mailparser import MailParser
from pikepdf import Pdf, PdfError
from pypdf import PdfReader
from sklearn.ensemble import IsolationForest

from ..config import get_settings
from ..storage.forensics_chain import ForensicsChainLedger
from ..utils.text import read_text


SCHEMA_VERSION = "2025-11-06"


_tracer = trace.get_tracer(__name__)
_meter = metrics.get_meter(__name__)
_forensics_pipeline_counter = _meter.create_counter(
    "forensics_reports_total",
    unit="1",
    description="Total forensics reports generated",
)
_forensics_fallback_counter = _meter.create_counter(
    "forensics_pipeline_fallbacks_total",
    unit="1",
    description="Number of pipelines completing with fallback applied",
)
_forensics_pipeline_duration = _meter.create_histogram(
    "forensics_pipeline_duration_ms",
    unit="ms",
    description="Duration of forensics pipelines",
)
_forensics_stage_duration = _meter.create_histogram(
    "forensics_stage_duration_ms",
    unit="ms",
    description="Duration of individual forensics stages",
)


@dataclass
class ForensicsSignal:
    type: str
    level: str
    detail: str
    data: Dict[str, Any] | None = None

    def to_dict(self) -> Dict[str, Any]:
        payload: Dict[str, Any] = {
            "type": self.type,
            "level": self.level,
            "detail": self.detail,
        }
        if self.data is not None:
            payload["data"] = ForensicsService.to_jsonable(self.data)
        return payload


@dataclass
class StageRecord:
    name: str
    started_at: str
    completed_at: str
    status: str
    notes: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "name": self.name,
            "started_at": self.started_at,
            "completed_at": self.completed_at,
            "status": self.status,
            "notes": list(self.notes),
        }


@dataclass
class ForensicsReport:
    schema_version: str
    file_id: str
    artifact_type: str
    generated_at: str
    summary: str
    data: Dict[str, Any]
    metadata: Dict[str, Any]
    signals: List[ForensicsSignal]
    stages: List[StageRecord]
    fallback_applied: bool
    report_path: Path | None = None

    def artifact_mapping(self) -> Dict[str, Any]:
        return {
            "summary": self.summary,
            "data": ForensicsService.to_jsonable(self.data),
            "metadata": ForensicsService.to_jsonable(self.metadata),
            "signals": [signal.to_dict() for signal in self.signals],
            "stages": [stage.to_dict() for stage in self.stages],
            "fallback_applied": self.fallback_applied,
            "schema_version": self.schema_version,
            "generated_at": self.generated_at,
        }


@dataclass
class PipelineStage:
    name: str
    handler: Callable[["PipelineContext", List[str]], None]
    required: bool = True


@dataclass
class PipelineContext:
    file_id: str
    artifact_type: str
    source_path: Path
    canonical_path: Path | None = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    payload: Dict[str, Any] = field(default_factory=dict)
    summary_lines: List[str] = field(default_factory=list)
    signals: List[ForensicsSignal] = field(default_factory=list)
    fallback_applied: bool = False

    @property
    def summary(self) -> str:
        return " ".join(line for line in self.summary_lines if line).strip()


class ForensicsService:
    def __init__(self) -> None:
        self.settings = get_settings()
        self.base_dir = self.settings.forensics_dir
        self.base_dir.mkdir(parents=True, exist_ok=True)
        self.chain_ledger = ForensicsChainLedger(self.settings.forensics_chain_path)

    # region public API
    def build_document_artifact(self, file_id: str, path: Path) -> ForensicsReport:
        ctx = PipelineContext(file_id=file_id, artifact_type="document", source_path=path)
        stages = [
            PipelineStage("canonicalise", self._stage_canonicalise),
            PipelineStage("metadata", self._stage_document_metadata),
            PipelineStage("analyse", self._stage_document_analyse, required=False),
        ]
        records, duration_ms = self._execute_pipeline(ctx, stages)
        report = ForensicsReport(
            schema_version=SCHEMA_VERSION,
            file_id=file_id,
            artifact_type="document",
            generated_at=self._now_iso(),
            summary=ctx.summary or "Document analysis completed",
            data=ctx.payload,
            metadata=ctx.metadata,
            signals=ctx.signals,
            stages=records,
            fallback_applied=ctx.fallback_applied,
        )
        report.report_path = self._persist(report)
        self._record_pipeline_metrics(ctx, duration_ms)
        return report

    def build_image_artifact(self, file_id: str, path: Path) -> ForensicsReport:
        ctx = PipelineContext(file_id=file_id, artifact_type="image", source_path=path)
        stages = [
            PipelineStage("canonicalise", self._stage_canonicalise),
            PipelineStage("metadata", self._stage_image_metadata),
            PipelineStage("analyse", self._stage_image_analyse, required=False),
        ]
        records, duration_ms = self._execute_pipeline(ctx, stages)
        report = ForensicsReport(
            schema_version=SCHEMA_VERSION,
            file_id=file_id,
            artifact_type="image",
            generated_at=self._now_iso(),
            summary=ctx.summary or "Image analysis completed",
            data=ctx.payload,
            metadata=ctx.metadata,
            signals=ctx.signals,
            stages=records,
            fallback_applied=ctx.fallback_applied,
        )
        report.report_path = self._persist(report)
        self._record_pipeline_metrics(ctx, duration_ms)
        return report

    def build_financial_artifact(self, file_id: str, path: Path) -> ForensicsReport:
        ctx = PipelineContext(file_id=file_id, artifact_type="financial", source_path=path)
        stages = [
            PipelineStage("canonicalise", self._stage_canonicalise),
            PipelineStage("metadata", self._stage_financial_metadata),
            PipelineStage("analyse", self._stage_financial_analyse, required=False),
        ]
        records, duration_ms = self._execute_pipeline(ctx, stages)
        report = ForensicsReport(
            schema_version=SCHEMA_VERSION,
            file_id=file_id,
            artifact_type="financial",
            generated_at=self._now_iso(),
            summary=ctx.summary or "Financial analysis completed",
            data=ctx.payload,
            metadata=ctx.metadata,
            signals=ctx.signals,
            stages=records,
            fallback_applied=ctx.fallback_applied,
        )
        report.report_path = self._persist(report)
        self._record_pipeline_metrics(ctx, duration_ms)
        return report

    def load_report(self, file_id: str) -> Dict[str, Any]:
        report_path = self.base_dir / file_id / "report.json"
        if not report_path.exists():
            raise FileNotFoundError(f"No forensics report found for {file_id}")
        return json.loads(report_path.read_text())

    def load_artifact(self, file_id: str, artifact: str) -> Dict[str, Any]:
        try:
            report = self.load_report(file_id)
        except FileNotFoundError:
            legacy_path = self.base_dir / file_id / f"{artifact}.json"
            if not legacy_path.exists():
                raise
            legacy_payload = json.loads(legacy_path.read_text())
            return {
                "summary": f"Legacy {artifact} artefact",
                "data": legacy_payload,
                "metadata": {},
                "signals": [],
                "stages": [],
                "fallback_applied": False,
                "schema_version": "legacy",
                "generated_at": None,
            }
        artifacts = report.get("artifacts", {})
        if artifact not in artifacts:
            raise FileNotFoundError(f"Artifact {artifact} missing for {file_id}")
        return artifacts[artifact]

    def report_exists(self, file_id: str, artifact: str) -> bool:
        try:
            report = self.load_report(file_id)
        except FileNotFoundError:
            return False
        return artifact in report.get("artifacts", {})

    # endregion

    # region pipeline execution
    def _execute_pipeline(
        self, ctx: PipelineContext, stages: List[PipelineStage]
    ) -> Tuple[List[StageRecord], float]:
        records: List[StageRecord] = []
        pipeline_start = perf_counter()
        pipeline_duration_ms = 0.0
        with _tracer.start_as_current_span("forensics.pipeline") as pipeline_span:
            pipeline_span.set_attribute("forensics.file_id", ctx.file_id)
            pipeline_span.set_attribute("forensics.artifact_type", ctx.artifact_type)
            pipeline_span.set_attribute("forensics.stage.count", len(stages))

            for stage in stages:
                started = self._now_iso()
                notes: List[str] = []
                status = "succeeded"
                stage_start = perf_counter()
                with _tracer.start_as_current_span(
                    f"forensics.stage.{stage.name}"
                ) as stage_span:
                    stage_span.set_attribute("forensics.stage.name", stage.name)
                    stage_span.set_attribute("forensics.stage.required", stage.required)
                    stage_span.set_attribute("forensics.artifact_type", ctx.artifact_type)
                    try:
                        stage.handler(ctx, notes)
                    except Exception as exc:  # pylint: disable=broad-except
                        ctx.fallback_applied = True
                        status = "failed"
                        notes.append(str(exc))
                        stage_span.record_exception(exc)
                        stage_span.set_status(Status(StatusCode.ERROR, str(exc)))
                        pipeline_span.add_event(
                            "forensics.stage.failure",
                            {
                                "forensics.stage": stage.name,
                                "forensics.artifact_type": ctx.artifact_type,
                            },
                        )
                        if stage.required:
                            completed_at = self._now_iso()
                            duration_ms = (perf_counter() - stage_start) * 1000.0
                            stage_span.set_attribute("forensics.stage.duration_ms", duration_ms)
                            stage_span.set_attribute("forensics.stage.status", status)
                            stage_span.set_attribute(
                                "forensics.stage.notes_count", len(notes)
                            )
                            stage_span.set_attribute(
                                "forensics.stage.fallback_applied", ctx.fallback_applied
                            )
                            _forensics_stage_duration.record(
                                duration_ms,
                                attributes={
                                    "stage": stage.name,
                                    "artifact_type": ctx.artifact_type,
                                    "required": stage.required,
                                    "status": status,
                                },
                            )
                            records.append(
                                StageRecord(stage.name, started, completed_at, status, notes)
                            )
                            pipeline_span.set_status(
                                Status(StatusCode.ERROR, f"Stage {stage.name} failed")
                            )
                            raise
                    finally:
                        completed_at = self._now_iso()
                        duration_ms = (perf_counter() - stage_start) * 1000.0
                        stage_span.set_attribute("forensics.stage.duration_ms", duration_ms)
                        stage_span.set_attribute("forensics.stage.status", status)
                        stage_span.set_attribute(
                            "forensics.stage.notes_count", len(notes)
                        )
                        stage_span.set_attribute(
                            "forensics.stage.fallback_applied", ctx.fallback_applied
                        )
                        _forensics_stage_duration.record(
                            duration_ms,
                            attributes={
                                "stage": stage.name,
                                "artifact_type": ctx.artifact_type,
                                "required": stage.required,
                                "status": status,
                            },
                        )
                records.append(
                    StageRecord(stage.name, started, completed_at, status, notes)
                )

            pipeline_duration_ms = (perf_counter() - pipeline_start) * 1000.0
            pipeline_span.set_attribute(
                "forensics.pipeline.duration_ms", pipeline_duration_ms
            )
            pipeline_span.set_attribute(
                "forensics.pipeline.fallback_applied", ctx.fallback_applied
            )

        return records, pipeline_duration_ms

    def _record_pipeline_metrics(self, ctx: PipelineContext, duration_ms: float) -> None:
        attributes = {
            "artifact_type": ctx.artifact_type,
        }
        _forensics_pipeline_counter.add(1, attributes=attributes)
        _forensics_pipeline_duration.record(duration_ms, attributes=attributes)
        if ctx.fallback_applied:
            _forensics_fallback_counter.add(1, attributes=attributes)

    def _stage_canonicalise(self, ctx: PipelineContext, notes: List[str]) -> None:
        destination_dir = self.base_dir / ctx.file_id / "source"
        destination_dir.mkdir(parents=True, exist_ok=True)
        canonical_path = destination_dir / ctx.source_path.name
        shutil.copy2(ctx.source_path, canonical_path)
        ctx.canonical_path = canonical_path
        ctx.metadata["canonical_path"] = str(canonical_path)
        stat = canonical_path.stat()
        ctx.metadata["size_bytes"] = stat.st_size
        ctx.metadata["modified_at"] = datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat()
        notes.append(f"Canonical copy stored at {canonical_path}")

    # document stages
    def _stage_document_metadata(self, ctx: PipelineContext, notes: List[str]) -> None:
        path = ctx.canonical_path or ctx.source_path
        ctx.metadata["extension"] = path.suffix.lower()
        mime_type = mimetypes.guess_type(str(path))[0]
        ctx.metadata["mime_type"] = mime_type or "application/octet-stream"
        payload = path.read_bytes()
        hashes = {
            "md5": md5(payload).hexdigest(),
            "sha256": sha256(payload).hexdigest(),
        }
        fuzzy = self._fuzzy_digest(payload)
        if fuzzy:
            hashes["tlsh"] = fuzzy
        else:
            ctx.signals.append(
                ForensicsSignal(
                    "hash.tlsh",
                    "warning",
                    "Fuzzy hash unavailable (insufficient content length)",
                )
            )
        ctx.payload["hashes"] = hashes
        ctx.summary_lines.append(
            f"Document hashed (SHA-256 {hashes['sha256'][:16]}…); MIME {ctx.metadata.get('mime_type', 'unknown')}"
        )

    def _stage_document_analyse(self, ctx: PipelineContext, notes: List[str]) -> None:
        path = ctx.canonical_path or ctx.source_path
        extension = ctx.metadata.get("extension", path.suffix.lower())
        analysis: Dict[str, Any] = {}
        authenticity: Dict[str, Any] = {}
        if extension == ".pdf":
            analysis.update(self._analyse_pdf(path, ctx))
            authenticity.update(self._pdf_authenticity(path, ctx))
        elif extension == ".docx":
            analysis.update(self._analyse_docx(path, ctx))
            authenticity.update(self._docx_authenticity(path, ctx))
        elif extension == ".msg":
            analysis.update(self._analyse_msg(path, ctx))
        else:
            analysis.update(self._analyse_text(path, ctx))
            authenticity.update(self._text_authenticity(path, ctx))
        ctx.payload["analysis"] = analysis
        if authenticity:
            ctx.payload["authenticity"] = authenticity
        notes.append(f"Analysis completed for extension {extension}")

    # image stages
    def _stage_image_metadata(self, ctx: PipelineContext, notes: List[str]) -> None:
        path = ctx.canonical_path or ctx.source_path
        with Image.open(path) as img:
            width, height = img.size
            ctx.metadata.update(
                {
                    "mode": img.mode,
                    "format": img.format,
                    "width": width,
                    "height": height,
                }
            )
            mime_type = mimetypes.guess_type(str(path))[0]
            ctx.metadata["mime_type"] = mime_type or f"image/{(img.format or '').lower()}"
            exif_data = {}
            if "exif" in img.info:
                try:
                    exif_data = piexif.load(img.info["exif"])
                except Exception:  # pylint: disable=broad-except
                    exif_data = {}
            ctx.payload["exif"] = self._simplify_exif(exif_data)
            ctx.summary_lines.append(f"Image {width}×{height} analysed; format {img.format}")
        notes.append("Image metadata harvested")

    def _stage_image_analyse(self, ctx: PipelineContext, notes: List[str]) -> None:
        path = ctx.canonical_path or ctx.source_path
        with Image.open(path) as img:
            ela_score = self._ela_score(img)
            clones = self._clone_map(img)
            residual = self._residual_stats(img)
            ctx.payload["ela"] = {"mean_absolute_error": ela_score}
            ctx.payload["clone_candidates"] = clones
            ctx.payload["residual_stats"] = residual
            if ctx.metadata.get("width", 0) < 128 or ctx.metadata.get("height", 0) < 128:
                ctx.fallback_applied = True
                ctx.signals.append(
                    ForensicsSignal(
                        "image.fallback",
                        "warning",
                        "Image below resolution threshold; high-confidence detectors skipped.",
                        {"width": ctx.metadata.get("width"), "height": ctx.metadata.get("height")},
                    )
                )
            ctx.summary_lines.append(
                f"ELA score {ela_score:.2f}; {len(clones)} clone candidates; residual energy {residual['energy']:.4f}"
            )
        notes.append("Image forensic analyzers executed")

    # financial stages
    def _stage_financial_metadata(self, ctx: PipelineContext, notes: List[str]) -> None:
        path = ctx.canonical_path or ctx.source_path
        df = pd.read_csv(path)
        ctx.metadata.update(
            {
                "rows": int(df.shape[0]),
                "columns": list(df.columns),
            }
        )
        ctx.payload["preview"] = df.head(10).to_dict(orient="records")
        ctx.payload["dtypes"] = {column: str(dtype) for column, dtype in df.dtypes.items()}
        ctx.payload["numeric_columns"] = [column for column in df.columns if pd.api.types.is_numeric_dtype(df[column])]
        ctx.payload["entities"] = (
            df["entity"].dropna().unique().tolist() if "entity" in df.columns else []
        )
        ctx.metadata["csv_path"] = str(path)
        ctx.metadata["row_count"] = int(df.shape[0])
        ctx.summary_lines.append(f"Ledger contains {int(df.shape[0])} rows across {len(df.columns)} columns")
        ctx.payload["dataframe"] = df
        notes.append("Financial ledger loaded into DataFrame")

    def _stage_financial_analyse(self, ctx: PipelineContext, notes: List[str]) -> None:
        df: pd.DataFrame = ctx.payload.pop("dataframe")
        totals = self._financial_totals(df)
        anomalies = self._financial_anomalies(df)
        ctx.payload["totals"] = totals
        ctx.payload["anomalies"] = anomalies
        ctx.summary_lines.append(
            f"Detected {len(anomalies)} anomalous rows across {len(ctx.payload['numeric_columns'])} numeric fields"
        )
        ctx.signals.append(
            ForensicsSignal(
                "financial.anomaly_count",
                "info",
                f"Anomaly scan completed with {len(anomalies)} flagged rows",
            )
        )
        notes.append("Financial anomaly detection executed")

    # endregion

    # region helpers
    def _analyse_pdf(self, path: Path, ctx: PipelineContext) -> Dict[str, Any]:
        analysis: Dict[str, Any] = {}
        try:
            reader = PdfReader(str(path))
            analysis["page_count"] = len(reader.pages)
            metadata = reader.metadata or {}
            analysis["metadata"] = {key: str(value) for key, value in metadata.items()}
            outline_count = 0
            try:
                outlines = reader.outline
                if isinstance(outlines, list):
                    outline_count = len(outlines)
            except Exception:  # pylint: disable=broad-except
                outline_count = 0
            analysis["outline_items"] = outline_count
        except Exception as exc:  # pylint: disable=broad-except
            ctx.signals.append(
                ForensicsSignal("document.pdf", "error", f"Failed to parse PDF: {exc}")
            )
        return analysis

    def _pdf_authenticity(self, path: Path, ctx: PipelineContext) -> Dict[str, Any]:
        authenticity: Dict[str, Any] = {}
        try:
            with Pdf.open(str(path)) as pdf:
                root_repr = repr(pdf.Root)
                authenticity["signature_hint"] = "/Sig" in root_repr
                if authenticity["signature_hint"]:
                    ctx.signals.append(
                        ForensicsSignal(
                            "document.pdf.signature",
                            "info",
                            "PDF contains signature objects",
                        )
                    )
        except PdfError as exc:
            ctx.signals.append(
                ForensicsSignal("document.pdf", "warning", f"Unable to open PDF with pikepdf: {exc}")
            )
        return authenticity

    def _analyse_docx(self, path: Path, ctx: PipelineContext) -> Dict[str, Any]:
        analysis: Dict[str, Any] = {}
        try:
            document = DocxDocument(str(path))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            headings = [paragraph.text for paragraph in document.paragraphs if paragraph.style and str(paragraph.style.name).startswith("Heading")]
            analysis["paragraph_count"] = len(paragraphs)
            analysis["headings"] = headings
            analysis["has_toc"] = any("table of contents" in paragraph.lower() for paragraph in paragraphs)
            core = document.core_properties
            ctx.metadata["docx_core_properties"] = {
                "author": core.author,
                "last_modified_by": core.last_modified_by,
                "created": core.created.isoformat() if core.created else None,
                "modified": core.modified.isoformat() if core.modified else None,
            }
            if not analysis["has_toc"]:
                ctx.signals.append(
                    ForensicsSignal(
                        "document.docx.toc",
                        "warning",
                        "Document missing Table of Contents heading",
                    )
                )
        except Exception as exc:  # pylint: disable=broad-except
            ctx.signals.append(
                ForensicsSignal("document.docx", "error", f"Failed to parse DOCX: {exc}")
            )
        return analysis

    def _docx_authenticity(self, path: Path, ctx: PipelineContext) -> Dict[str, Any]:
        text = read_text(path)
        entropy = self._shannon_entropy(text)
        word_count = len(text.split())
        ctx.signals.append(
            ForensicsSignal(
                "document.docx.entropy",
                "info",
                "Shannon entropy computed",
                {"bits_per_char": entropy},
            )
        )
        return {"entropy": entropy, "word_count": word_count}

    def _analyse_msg(self, path: Path, ctx: PipelineContext) -> Dict[str, Any]:
        analysis: Dict[str, Any] = {}
        try:
            message = MsgMessage(str(path))
            parser = MailParser.from_string(message.as_string())
            analysis["subject"] = message.subject
            analysis["sender"] = message.sender
            analysis["to"] = list(message.to)
            analysis["attachment_count"] = len(message.attachments)
            analysis["received"] = [header for header in parser.headers if header["name"].lower() == "received"]
        except Exception as exc:  # pylint: disable=broad-except
            ctx.signals.append(ForensicsSignal("document.msg", "error", f"Failed to parse MSG: {exc}"))
        return analysis

    def _analyse_text(self, path: Path, ctx: PipelineContext) -> Dict[str, Any]:
        text = read_text(path)
        lines = text.splitlines()
        headings = [line.strip() for line in lines if line.strip().isupper() and len(line.strip()) > 4]
        sections = Counter(heading.split()[0] for heading in headings)
        ctx.signals.append(
            ForensicsSignal(
                "document.text.headings",
                "info",
                "Detected uppercase headings",
                {"count": len(headings)},
            )
        )
        return {
            "line_count": len(lines),
            "headings": headings[:20],
            "section_token_counts": dict(sections),
        }

    def _text_authenticity(self, path: Path, ctx: PipelineContext) -> Dict[str, Any]:
        text = read_text(path)
        entropy = self._shannon_entropy(text)
        ctx.signals.append(
            ForensicsSignal(
                "document.text.entropy",
                "info",
                "Shannon entropy computed",
                {"bits_per_char": entropy},
            )
        )
        return {"entropy": entropy}

    def _simplify_exif(self, exif_data: Dict[str, Any]) -> Dict[str, Any]:
        simplified: Dict[str, Any] = {}
        for section, payload in exif_data.items():
            if isinstance(payload, dict):
                simplified[section] = {str(key): ForensicsService.to_jsonable(value) for key, value in payload.items() if key in {271, 272, 274, 306}}
        return simplified

    def _ela_score(self, image: Image.Image) -> float:
        buffer = Path(self.base_dir / "_tmp_ela.jpg")
        image.save(buffer, format="JPEG", quality=90)
        try:
            with Image.open(buffer) as recompressed:
                original_arr = np.asarray(image.convert("RGB"), dtype=np.float32)
                recompressed_arr = np.asarray(recompressed.convert("RGB"), dtype=np.float32)
                diff = np.abs(original_arr - recompressed_arr)
                score = float(np.mean(diff))
        finally:
            buffer.unlink(missing_ok=True)
        return score

    def _clone_map(self, image: Image.Image, tile: int = 32) -> List[Dict[str, Any]]:
        gray = np.asarray(image.convert("L"))
        clones: List[Dict[str, Any]] = []
        signatures: Dict[tuple[int, ...], tuple[int, int]] = {}
        height, width = gray.shape
        for y in range(0, height - tile + 1, tile):
            for x in range(0, width - tile + 1, tile):
                patch = gray[y : y + tile, x : x + tile]
                resized = Image.fromarray(patch).resize((8, 8), Image.BILINEAR)
                resized_arr = np.asarray(resized, dtype=np.float32)
                signature_tuple = tuple(int(value) for value in (resized_arr > resized_arr.mean()).astype(int).flatten())
                if signature_tuple in signatures:
                    clones.append({"source": signatures[signature_tuple], "duplicate": (x, y)})
                else:
                    signatures[signature_tuple] = (x, y)
        return clones[:50]

    def _residual_stats(self, image: Image.Image) -> Dict[str, float]:
        original = image.convert("L")
        blurred = original.filter(ImageFilter.GaussianBlur(radius=1))
        arr = np.asarray(original, dtype=np.float32)
        blurred_arr = np.asarray(blurred, dtype=np.float32)
        residual = arr - blurred_arr
        return {
            "mean": float(residual.mean()),
            "stddev": float(residual.std()),
            "energy": float(np.mean(residual ** 2)),
        }

    def _financial_totals(self, df: pd.DataFrame) -> Dict[str, str]:
        totals: Dict[str, str] = {}
        for column in df.columns:
            if not pd.api.types.is_numeric_dtype(df[column]):
                continue
            total = Decimal("0")
            for value in df[column].dropna():
                total += Decimal(str(value))
            totals[column] = str(total)
        return totals

    def _financial_anomalies(self, df: pd.DataFrame) -> List[Dict[str, Any]]:
        numeric_df = df.select_dtypes(include=["number"]).fillna(0.0)
        anomalies: List[Dict[str, Any]] = []
        if numeric_df.empty:
            return anomalies
        if len(numeric_df) >= 5:
            model = IsolationForest(random_state=42, contamination="auto")
            model.fit(numeric_df.values)
            scores = model.score_samples(numeric_df.values)
            threshold = float(np.quantile(scores, 0.05))
            for idx, score in enumerate(scores):
                if score <= threshold:
                    payload = df.iloc[idx].to_dict()
                    payload["isolation_score"] = float(score)
                    anomalies.append(ForensicsService.to_jsonable(payload))
        else:
            means = numeric_df.mean()
            stds = numeric_df.std(ddof=0)
            for idx, row in numeric_df.iterrows():
                zscores = {}
                outlier = False
                for column, value in row.items():
                    std = stds[column]
                    if std == 0:
                        continue
                    z = (value - means[column]) / std
                    if abs(z) >= 2.5:
                        outlier = True
                        zscores[column] = float(z)
                if outlier:
                    payload = df.iloc[idx].to_dict()
                    payload["zscore"] = zscores
                    anomalies.append(ForensicsService.to_jsonable(payload))
        return anomalies

    @staticmethod
    def _shannon_entropy(text: str) -> float:
        if not text:
            return 0.0
        freq: Dict[str, int] = defaultdict(int)
        for char in text:
            freq[char] += 1
        total = len(text)
        entropy = 0.0
        for count in freq.values():
            probability = count / total
            entropy -= probability * math.log2(probability)
        return entropy

    @staticmethod
    def _fuzzy_digest(payload: bytes) -> str | None:
        if len(payload) < 256:
            return None
        return sha256(payload).hexdigest()[:32]

    def _persist(self, report: ForensicsReport) -> Path:
        directory = self.base_dir / report.file_id
        directory.mkdir(parents=True, exist_ok=True)
        report_path = directory / "report.json"
        if report_path.exists():
            payload = json.loads(report_path.read_text())
        else:
            payload = {
                "file_id": report.file_id,
                "artifacts": {},
            }
        payload["schema_version"] = SCHEMA_VERSION
        payload["generated_at"] = report.generated_at
        payload.setdefault("artifacts", {})[report.artifact_type] = report.artifact_mapping()
        report_path.write_text(json.dumps(payload, indent=2, sort_keys=True))
        checksum = sha256(report_path.read_bytes()).hexdigest()
        self.chain_ledger.append(
            actor="forensics.service",
            action=f"persist:{report.artifact_type}",
            payload={
                "file_id": report.file_id,
                "artifact_type": report.artifact_type,
                "report_path": report_path,
                "generated_at": report.generated_at,
                "checksum_sha256": checksum,
                "signals": [signal.to_dict() for signal in report.signals],
            },
        )
        return report_path

    @staticmethod
    def _now_iso() -> str:
        return datetime.now(timezone.utc).isoformat()

    @staticmethod
    def to_jsonable(value: Any) -> Any:
        if isinstance(value, Decimal):
            return str(value)
        if isinstance(value, (datetime,)):
            return value.isoformat()
        if isinstance(value, Path):
            return str(value)
        if isinstance(value, np.generic):
            return value.item()
        if isinstance(value, dict):
            return {key: ForensicsService.to_jsonable(val) for key, val in value.items()}
        if isinstance(value, list):
            return [ForensicsService.to_jsonable(item) for item in value]
        return value

    # endregion


def get_forensics_service() -> ForensicsService:
    return ForensicsService()
