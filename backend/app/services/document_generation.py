from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt
import os

class DocumentGenerationService:
    def draft_legal_document(self, filepath: str, motion_type: str, data: Dict[str, Any]) -> str:
        """
        Drafts a legal document based on the motion type and data.
        """
        doc = Document()
        
        # Title
        title = doc.add_heading(f"MOTION FOR {motion_type.upper()}", 0)
        title.alignment = 1 # Center
        
        # Case Caption (simplified)
        doc.add_paragraph(f"CASE ID: {data.get('case_id', 'UNKNOWN')}")
        doc.add_paragraph(f"DATE: {data.get('date', 'UNKNOWN')}")
        
        doc.add_heading('INTRODUCTION', level=1)
        doc.add_paragraph(data.get('introduction', 'No introduction provided.'))
        
        doc.add_heading('FACTS', level=1)
        facts = data.get('facts', [])
        if isinstance(facts, list):
            for fact in facts:
                doc.add_paragraph(str(fact), style='List Bullet')
        else:
            doc.add_paragraph(str(facts))
            
        doc.add_heading('LEGAL ARGUMENT', level=1)
        doc.add_paragraph(data.get('argument', 'No argument provided.'))
        
        doc.add_heading('CONCLUSION', level=1)
        doc.add_paragraph(data.get('conclusion', 'No conclusion provided.'))
        
        # Save
        doc.save(filepath)
        return filepath

    def prepare_binder(self, filepath: str, evidence_list: List[Dict[str, Any]], case_name: str) -> str:
        """
        Prepares a trial binder document.
        """
        doc = Document()
        
        title = doc.add_heading(f"TRIAL BINDER: {case_name}", 0)
        title.alignment = 1
        
        doc.add_paragraph("Table of Contents")
        
        for i, item in enumerate(evidence_list, 1):
            doc.add_heading(f"Exhibit {i}: {item.get('name', 'Untitled')}", level=1)
            
            p = doc.add_paragraph()
            p.add_run("ID: ").bold = True
            p.add_run(str(item.get('id', 'N/A')))
            
            p = doc.add_paragraph()
            p.add_run("Type: ").bold = True
            p.add_run(str(item.get('type', 'Unknown')))
            
            p = doc.add_paragraph()
            p.add_run("URL: ").bold = True
            p.add_run(str(item.get('url', 'N/A')))
            
            if item.get('annotation'):
                doc.add_heading("Annotation", level=2)
                doc.add_paragraph(item['annotation'])
                
            doc.add_page_break()
            
        doc.save(filepath)
        return filepath

_service = DocumentGenerationService()

def get_document_generation_service() -> DocumentGenerationService:
    return _service
